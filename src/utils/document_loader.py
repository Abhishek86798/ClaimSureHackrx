"""
Document loader utility functions for Claimsure.

Provides functions to load and clean text from various document formats.
"""

import os
import re
import io
import logging
from typing import Optional
from pathlib import Path
from urllib.parse import urlparse
import requests

import PyPDF2
from docx import Document
import email
from email import policy

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing extra spaces, headers, footers, and contact info.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text as a single string
    """
    if not text:
        return ""
    
    # Remove common headers/footers first
    header_patterns = [
        r'Page \d+ of \d+',
        r'Page \d+',
        r'Confidential',
        r'Draft',
        r'Internal Use Only',
        r'© \d{4}.*?All rights reserved\.',
        r'Copyright © \d{4}.*?',
    ]
    
    for pattern in header_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Remove email signatures (only at the end of text)
    signature_patterns = [
        r'--\s*\n.*$',  # Email signatures after -- (end of text)
    ]
    
    for pattern in signature_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove signature and contact lines (only standalone lines)
    lines_to_remove = [
        r'^\s*Best regards,.*$',
        r'^\s*Sincerely,.*$',
        r'^\s*Thank you,.*$',
        r'^\s*Phone:.*$',
        r'^\s*Email:.*$',
        r'^\s*www\..*$',
        r'^\s*http[s]?://.*$',
    ]
    
    for pattern in lines_to_remove:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove multiple newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]
    
    # Join lines and normalize spaces
    result = ' '.join(lines)
    result = re.sub(r'\s+', ' ', result)
    
    return result.strip()


def load_pdf(path_or_url: str) -> str:
    """
    Extract clean text from PDF files or URLs.
    
    Args:
        path_or_url: File path or URL to PDF
        
    Returns:
        Cleaned text as a single string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If URL is invalid or file can't be loaded
    """
    try:
        # Check if it's a URL
        parsed = urlparse(path_or_url)
        if parsed.scheme and parsed.netloc:
            # Download from URL
            response = requests.get(path_or_url, timeout=30)
            response.raise_for_status()
            pdf_content = response.content
        else:
            # Load from local file
            file_path = Path(path_or_url)
            if not file_path.exists():
                raise FileNotFoundError(f"PDF file not found: {path_or_url}")
            
            with open(file_path, 'rb') as file:
                pdf_content = file.read()
        
        # Extract text from PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        
        content = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                content += page_text + "\n"
        
        return clean_text(content)
        
    except Exception as e:
        logger.error(f"Error loading PDF {path_or_url}: {str(e)}")
        raise ValueError(f"Failed to load PDF: {str(e)}")


def load_docx(path: str) -> str:
    """
    Extract clean text from DOCX files.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        Cleaned text as a single string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file can't be loaded
    """
    try:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        doc = Document(file_path)
        
        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content += cell.text + "\n"
        
        return clean_text(content)
        
    except Exception as e:
        logger.error(f"Error loading DOCX {path}: {str(e)}")
        raise ValueError(f"Failed to load DOCX: {str(e)}")


def load_email(path: str) -> str:
    """
    Extract clean text from .eml files.
    
    Args:
        path: Path to .eml file
        
    Returns:
        Cleaned text as a single string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file can't be loaded
    """
    try:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"Email file not found: {path}")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            email_content = file.read()
        
        # Parse email
        msg = email.message_from_string(email_content, policy=policy.default)
        
        content = ""
        
        # Extract subject
        subject = msg.get('subject', '')
        if subject:
            content += f"Subject: {subject}\n\n"
        
        # Extract body
        if msg.is_multipart():
            # Handle multipart messages
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    try:
                        body = part.get_content()
                        if body:
                            content += body + "\n"
                    except Exception:
                        continue
        else:
            # Handle simple text messages
            body = msg.get_content()
            if body:
                content += body + "\n"
        
        return clean_text(content)
        
    except Exception as e:
        logger.error(f"Error loading email {path}: {str(e)}")
        raise ValueError(f"Failed to load email: {str(e)}")


def load_document(path_or_url: str) -> str:
    """
    Load and extract clean text from various document formats.
    
    Args:
        path_or_url: Path or URL to document
        
    Returns:
        Cleaned text as a single string
        
    Raises:
        ValueError: If file type is not supported or can't be loaded
    """
    path_or_url = str(path_or_url).lower()
    
    if path_or_url.endswith('.pdf'):
        return load_pdf(path_or_url)
    elif path_or_url.endswith('.docx'):
        return load_docx(path_or_url)
    elif path_or_url.endswith('.eml'):
        return load_email(path_or_url)
    elif path_or_url.endswith('.txt'):
        # Handle text files
        try:
            file_path = Path(path_or_url)
            if not file_path.exists():
                raise FileNotFoundError(f"Text file not found: {path_or_url}")
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return clean_text(content)
        except Exception as e:
            logger.error(f"Error loading text file {path_or_url}: {str(e)}")
            raise ValueError(f"Failed to load text file: {str(e)}")
    else:
        raise ValueError(f"Unsupported file type: {path_or_url}")


def validate_file_path(path: str) -> bool:
    """
    Validate if a file path exists and is accessible.
    
    Args:
        path: File path to validate
        
    Returns:
        True if file exists and is accessible
    """
    try:
        file_path = Path(path)
        return file_path.exists() and file_path.is_file()
    except Exception:
        return False


def get_file_size(path: str) -> int:
    """
    Get file size in bytes.
    
    Args:
        path: File path
        
    Returns:
        File size in bytes
        
    Raises:
        FileNotFoundError: If file doesn't exist
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    return file_path.stat().st_size
