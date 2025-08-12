"""
Document loader module for Claimsure.

Handles loading and parsing of various document formats including PDF, DOCX, and TXT files.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

import PyPDF2
from docx import Document
from config import SUPPORTED_FILE_TYPES, MAX_FILE_SIZE_BYTES

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Handles loading and parsing of documents from various file formats."""
    
    def __init__(self):
        self.supported_types = SUPPORTED_FILE_TYPES
        self.max_file_size = MAX_FILE_SIZE_BYTES
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from file and extract its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata and content
            
        Raises:
            ValueError: If file type is not supported or file is too large
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File too large: {file_path.stat().st_size} bytes")
        
        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            if file_extension == ".pdf":
                return self._load_pdf(file_path)
            elif file_extension == ".docx":
                return self._load_docx(file_path)
            elif file_extension == ".txt":
                return self._load_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Load and parse PDF document."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                return {
                    "file_path": str(file_path),
                    "file_type": "pdf",
                    "content": content,
                    "metadata": {
                        "num_pages": len(pdf_reader.pages),
                        "file_size": file_path.stat().st_size,
                        "title": file_path.stem
                    }
                }
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """Load and parse DOCX document."""
        try:
            doc = Document(file_path)
            
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            return {
                "file_path": str(file_path),
                "file_type": "docx",
                "content": content,
                "metadata": {
                    "num_paragraphs": len(doc.paragraphs),
                    "file_size": file_path.stat().st_size,
                    "title": file_path.stem
                }
            }
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    def _load_txt(self, file_path: Path) -> Dict[str, Any]:
        """Load and parse TXT document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "file_path": str(file_path),
                "file_type": "txt",
                "content": content,
                "metadata": {
                    "file_size": file_path.stat().st_size,
                    "title": file_path.stem
                }
            }
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise
    
    def load_multiple_documents(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Load multiple documents from a list of file paths.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of loaded document dictionaries
        """
        documents = []
        for file_path in file_paths:
            try:
                doc = self.load_document(file_path)
                documents.append(doc)
                logger.info(f"Successfully loaded: {file_path}")
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {str(e)}")
                continue
        
        return documents
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate if a file can be loaded.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            True if file is valid, False otherwise
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return False
            
            if file_path.stat().st_size > self.max_file_size:
                return False
            
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_types:
                return False
            
            return True
        except Exception:
            return False
